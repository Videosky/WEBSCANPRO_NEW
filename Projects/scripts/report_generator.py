"""
WebScanProd - Report Generator for Windows
Generates PDF, DOCX, and Markdown vulnerability reports
"""

import pandas as pd
from pathlib import Path
import logging
from datetime import datetime

# Configure logging first
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Try to import optional dependencies with proper error handling
REPORTLAB_AVAILABLE = False
DOCX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib import colors
    from reportlab.lib.units import inch
    REPORTLAB_AVAILABLE = True
    logger.info("ReportLab loaded successfully - PDF reports enabled")
except ImportError as e:
    logger.warning("ReportLab not available - PDF reports disabled")
    logger.debug(f"Import error: {e}")

try:
    import docx
    from docx.shared import Inches
    DOCX_AVAILABLE = True
    logger.info("python-docx loaded successfully - DOCX reports enabled")
except ImportError as e:
    logger.warning("python-docx not available - DOCX reports disabled")
    logger.debug(f"Import error: {e}")

class ReportGenerator:
    def __init__(self, base_path="."):
        self.base_path = Path(base_path)
        self.processed_path = self.base_path / "data" / "processed"
        self.reports_path = self.base_path / "reports"
        
        # Ensure reports directory exists
        self.reports_path.mkdir(parents=True, exist_ok=True)
        
    def load_findings(self):
        """Load processed findings from CSV"""
        findings_file = self.processed_path / "findings.csv"
        if not findings_file.exists():
            raise FileNotFoundError(f"Processed findings not found at {findings_file}")
        
        df = pd.read_csv(findings_file)
        return df
    
    def generate_severity_summary(self, df):
        """Generate comprehensive severity summary"""
        severity_counts = df['severity'].value_counts()
        total = len(df)
        
        summary = {
            'total': total,
            'critical': severity_counts.get('Critical', 0),
            'high': severity_counts.get('High', 0),
            'medium': severity_counts.get('Medium', 0),
            'low': severity_counts.get('Low', 0),
            'risk_score': (
                severity_counts.get('Critical', 0) * 10 +
                severity_counts.get('High', 0) * 5 +
                severity_counts.get('Medium', 0) * 2 +
                severity_counts.get('Low', 0) * 1
            )
        }
        
        return summary
    
    def generate_markdown_report(self, df, summary):
        """Generate Markdown summary report"""
        report_file = self.reports_path / "webscan_summary.md"
        
        try:
            with open(report_file, 'w', encoding='utf-8') as f:
                f.write("# WebScanProd Vulnerability Assessment Report\n\n")
                f.write(f"**Report Date:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
                
                # Executive Summary
                f.write("## Executive Summary\n\n")
                f.write(f"Total Vulnerabilities Identified: **{summary['total']}**\n\n")
                f.write("### Severity Distribution\n")
                f.write(f"- 🔴 Critical: {summary['critical']}\n")
                f.write(f"- 🟠 High: {summary['high']}\n")
                f.write(f"- 🟡 Medium: {summary['medium']}\n")
                f.write(f"- 🔵 Low: {summary['low']}\n\n")
                f.write(f"**Overall Risk Score:** {summary['risk_score']}/100\n\n")
                
                # Detailed Findings
                f.write("## Detailed Vulnerability Findings\n\n")
                
                # Group by severity
                for severity in ['Critical', 'High', 'Medium', 'Low']:
                    severity_findings = df[df['severity'] == severity]
                    if not severity_findings.empty:
                        f.write(f"### {severity} Severity Findings\n\n")
                        
                        for _, finding in severity_findings.iterrows():
                            f.write(f"#### {finding['vuln_name']} ({finding['vuln_id']})\n")
                            f.write(f"- **Affected URL:** {finding['affected_url']}\n")
                            f.write(f"- **Description:** {finding['description']}\n")
                            f.write(f"- **Impact:** {finding['impact']}\n")
                            f.write(f"- **Recommendation:** {finding['recommendation']}\n\n")
                
                # Mitigation Summary
                f.write("## Recommended Mitigation Strategy\n\n")
                f.write("1. **Immediate Action Required:** Address all Critical and High severity issues\n")
                f.write("2. **Short-term Goals:** Resolve Medium severity vulnerabilities\n")
                f.write("3. **Long-term Improvements:** Implement security controls for Low severity findings\n")
            
            logger.info(f"Markdown report generated: {report_file}")
            return report_file
            
        except Exception as e:
            logger.error(f"Error generating Markdown report: {str(e)}")
            return None
    
    def generate_pdf_report(self, df, summary):
        """Generate PDF vulnerability report"""
        if not REPORTLAB_AVAILABLE:
            logger.warning("PDF generation disabled - ReportLab not installed")
            return None
            
        report_file = self.reports_path / "vulnerability_report.pdf"
        
        try:
            doc = SimpleDocTemplate(str(report_file), pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Title
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                textColor=colors.darkblue
            )
            title = Paragraph("WebScanProd Vulnerability Assessment Report", title_style)
            story.append(title)
            
            # Report Date
            date_style = ParagraphStyle(
                'CustomDate',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.grey
            )
            date_text = Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", date_style)
            story.append(date_text)
            story.append(Spacer(1, 20))
            
            # Executive Summary
            story.append(Paragraph("Executive Summary", styles['Heading2']))
            summary_text = f"""
            Total Vulnerabilities Identified: <b>{summary['total']}</b><br/>
            Critical: <font color="red"><b>{summary['critical']}</b></font> | 
            High: <font color="orange"><b>{summary['high']}</b></font> | 
            Medium: <font color="yellow"><b>{summary['medium']}</b></font> | 
            Low: <font color="blue"><b>{summary['low']}</b></font><br/>
            Overall Risk Score: <b>{summary['risk_score']}/100</b>
            """
            story.append(Paragraph(summary_text, styles['Normal']))
            story.append(Spacer(1, 20))
            
            # Detailed Findings Table
            story.append(Paragraph("Detailed Vulnerability Findings", styles['Heading2']))
            
            # Prepare table data
            table_data = [['ID', 'Vulnerability', 'Severity', 'Affected URL', 'Description']]
            
            for _, row in df.iterrows():
                # Truncate long descriptions for table display
                description = row['description']
                if len(description) > 100:
                    description = description[:100] + '...'
                
                table_data.append([
                    row['vuln_id'],
                    row['vuln_name'],
                    row['severity'],
                    row['affected_url'],
                    description
                ])
            
            # Create table
            table = Table(table_data, colWidths=[0.8*inch, 1.2*inch, 0.8*inch, 1.5*inch, 2.2*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('FONTSIZE', (0, 1), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            story.append(table)
            story.append(Spacer(1, 20))
            
            # Recommendations
            story.append(Paragraph("Recommended Mitigation Steps", styles['Heading2']))
            
            recommendations = []
            for _, row in df.iterrows():
                rec_text = f"<b>{row['vuln_id']} - {row['vuln_name']}:</b> {row['recommendation']}"
                recommendations.append(rec_text)
            
            rec_paragraph = Paragraph("<br/>\n".join(recommendations), styles['Normal'])
            story.append(rec_paragraph)
            
            # Build PDF
            doc.build(story)
            logger.info(f"PDF report generated: {report_file}")
            return report_file
            
        except Exception as e:
            logger.error(f"Error generating PDF report: {str(e)}")
            return None
    
    def generate_docx_report(self, df, summary):
        """Generate DOCX vulnerability report"""
        if not DOCX_AVAILABLE:
            logger.warning("DOCX generation disabled - python-docx not installed")
            return None
            
        report_file = self.reports_path / "vulnerability_report.docx"
        
        try:
            doc = docx.Document()
            
            # Title
            title = doc.add_heading('WebScanProd Vulnerability Assessment Report', 0)
            title.alignment = 1  # Center alignment
            
            # Report Date
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph()
            
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            summary_table = doc.add_table(rows=6, cols=2)
            summary_table.style = 'Light Shading'
            
            summary_data = [
                ['Total Vulnerabilities', str(summary['total'])],
                ['Critical', str(summary['critical'])],
                ['High', str(summary['high'])],
                ['Medium', str(summary['medium'])],
                ['Low', str(summary['low'])],
                ['Overall Risk Score', f"{summary['risk_score']}/100"]
            ]
            
            for i, (label, value) in enumerate(summary_data):
                summary_table.cell(i, 0).text = label
                summary_table.cell(i, 1).text = value
            
            doc.add_paragraph()
            
            # Detailed Findings
            doc.add_heading('Detailed Vulnerability Findings', level=1)
            
            # Create findings table
            findings_table = doc.add_table(rows=1, cols=5)
            findings_table.style = 'Light List'
            
            # Header row
            header_cells = findings_table.rows[0].cells
            headers = ['ID', 'Vulnerability', 'Severity', 'Affected URL', 'Description']
            for i, header in enumerate(headers):
                header_cells[i].text = header
            
            # Add findings rows
            for _, row in df.iterrows():
                row_cells = findings_table.add_row().cells
                row_cells[0].text = str(row['vuln_id'])
                row_cells[1].text = str(row['vuln_name'])
                row_cells[2].text = str(row['severity'])
                row_cells[3].text = str(row['affected_url'])
                row_cells[4].text = str(row['description'])
            
            doc.add_paragraph()
            
            # Recommendations
            doc.add_heading('Recommended Mitigation Steps', level=1)
            for _, row in df.iterrows():
                doc.add_heading(f"{row['vuln_id']} - {row['vuln_name']}", level=2)
                doc.add_paragraph(f"Recommendation: {row['recommendation']}")
                doc.add_paragraph(f"Impact: {row['impact']}")
                doc.add_paragraph()
            
            # Save document
            doc.save(str(report_file))
            logger.info(f"DOCX report generated: {report_file}")
            return report_file
            
        except Exception as e:
            logger.error(f"Error generating DOCX report: {str(e)}")
            return None
    
    def validate_reports(self, df):
        """Validate that all reports meet requirements"""
        try:
            validation_results = {
                'all_findings_have_description': not df['description'].isna().any() and (df['description'] != '').all(),
                'all_findings_have_severity': not df['severity'].isna().any() and (df['severity'] != '').all(),
                'all_findings_have_recommendation': not df['recommendation'].isna().any() and (df['recommendation'] != '').all(),
                'severity_levels_valid': df['severity'].isin(['Critical', 'High', 'Medium', 'Low']).all()
            }
            
            return all(validation_results.values()), validation_results
        except Exception as e:
            logger.error(f"Error during validation: {str(e)}")
            return False, {}
    
    def generate_all_reports(self):
        """Generate all report formats"""
        try:
            # Load findings
            df = self.load_findings()
            logger.info(f"Loaded {len(df)} vulnerabilities for reporting")
            
            # Generate summary
            summary = self.generate_severity_summary(df)
            
            # Validate data
            is_valid, validation_details = self.validate_reports(df)
            if not is_valid:
                logger.warning("Validation issues found:")
                for check, result in validation_details.items():
                    logger.warning(f"  {check}: {result}")
            else:
                logger.info("All data validation checks passed")
            
            # Generate reports
            reports_generated = []
            
            md_report = self.generate_markdown_report(df, summary)
            if md_report:
                reports_generated.append(("Markdown", md_report))
            
            pdf_report = self.generate_pdf_report(df, summary)
            if pdf_report:
                reports_generated.append(("PDF", pdf_report))
            
            docx_report = self.generate_docx_report(df, summary)
            if docx_report:
                reports_generated.append(("DOCX", docx_report))
            
            logger.info("Reports generated:")
            for report_type, report_path in reports_generated:
                logger.info(f"  - {report_type}: {report_path}")
            
            # At least Markdown should always be generated
            success = len(reports_generated) > 0
            if success:
                logger.info("WebScanProd reporting completed successfully!")
            else:
                logger.error("No reports were generated!")
            
            return success
            
        except Exception as e:
            logger.error(f"Error generating reports: {str(e)}")
            return False

def main():
    """Main execution function"""
    generator = ReportGenerator()
    success = generator.generate_all_reports()
    
    if success:
        logger.info("WebScanProd reporting completed successfully!")
    else:
        logger.error("WebScanProd reporting failed!")
    
    return success

if __name__ == "__main__":
    main()